import os
import argparse
from pathlib import Path

# Исправленный блок импорта
DOCX_AVAILABLE = False
try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    print("Предупреждение: Пакет python-docx не установлен. Используйте --format md.")

def collect_code(project_path, output_format='md'):
    project_path = Path(project_path).resolve()
    if not project_path.is_dir():
        raise ValueError("Указанный путь не является директорией.")

    ignore_dirs = {'venv', '.git', '__pycache__', 'migrations/versions', '.venv'}
    extensions = {'.py', '.md', '.txt', '.ini'}

    content = []
    file_count = 0

    for root, dirs, files in os.walk(project_path):
        dirs[:] = [d for d in dirs if d not in ignore_dirs]
        for file in files:
            if Path(file).suffix in extensions:
                file_path = Path(root) / file
                rel_path = file_path.relative_to(project_path)
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        code = f.read()
                    content.append(f"## Файл: {rel_path}\n\n```python\n{code}\n```\n\n")
                    file_count += 1
                    print(f"Обработан: {rel_path}")
                except Exception as e:
                    print(f"Ошибка чтения {rel_path}: {e}")

    if not content:
        raise ValueError("Нет файлов для сбора.")

    output_file = f"all_code.{output_format}"
    if output_format == 'md':
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write("# Сбор кода проекта\n\n" + ''.join(content))
    elif output_format == 'docx' and DOCX_AVAILABLE:
        doc = Document()
        doc.add_heading('Сбор кода проекта', 0)
        for section in content:
            doc.add_heading(section.split('\n')[0][3:], level=2)
            p = doc.add_paragraph(section.split('```python\n')[1].split('\n```')[0])
            p.style.font.name = 'Courier New'
            p.style.font.size = Pt(10)
        doc.save(output_file)
    else:
        raise ValueError("Формат 'docx' требует python-docx или используйте 'md'.")

    print(f"Готово! Создан файл: {output_file} ({file_count} файлов).")

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Сбор кода из проекта.')
    parser.add_argument('path', nargs='?', help='Путь к директории проекта')
    parser.add_argument('--format', default='md', choices=['md', 'docx'], help='Формат вывода: md или docx')
    args = parser.parse_args()

    path = args.path or input("Введите путь к директории проекта: ")
    try:
        collect_code(path, args.format)
    except Exception as e:
        print(f"Ошибка: {e}")